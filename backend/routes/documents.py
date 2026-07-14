"""Document routes — Upload, list, delete, move, QA, summarize, folders, tags, sharing."""
import os
import re
import json
import httpx
from datetime import datetime, timezone
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from models import (
    Document, DocumentChunk, DocumentQA, DocumentTag, DocumentShare,
    Folder, Tag, User
)
from services.autotag import auto_tag_document
from schemas import (
    DocumentResponse, DocumentMoveRequest, DocumentQARequest,
    DocumentQAResponse, DocumentSummarizeResponse,
    FolderCreate, FolderResponse, TagCreate, TagResponse, TagAssignRequest
)
from auth import get_current_user, get_db
from services.embeddings import get_index, rebuild_index

router = APIRouter(tags=["documents"])

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_TYPES = {".pdf", ".docx", ".txt", ".csv", ".md", ".doc"}


def _extract_text(filename: str, content: bytes) -> str:
    """Extract text from uploaded file."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt" or ext == ".md":
        return content.decode("utf-8", errors="ignore")

    if ext == ".csv":
        text = content.decode("utf-8", errors="ignore")
        lines = text.split("\n")
        return "\n".join(lines[:100])  # First 100 rows

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append(f"[Page {i+1}] {text}")
            return "\n\n".join(pages)
        except ImportError:
            return "[PDF processing requires pypdf]"

    if ext in (".docx", ".doc"):
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            return "[DOCX processing requires python-docx]"

    return content.decode("utf-8", errors="ignore")[:5000]


def _chunk_text(text: str, chunk_size: int = 1000) -> list[str]:
    """Split text into chunks."""
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    current = ""
    for sentence in sentences:
        if len(current) + len(sentence) > chunk_size:
            if current:
                chunks.append(current.strip())
            current = sentence
        else:
            current += " " + sentence if current else sentence
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _simple_search_score(query: str, text: str) -> float:
    """Simple keyword-based relevance score."""
    query_words = set(query.lower().split())
    text_lower = text.lower()
    matches = sum(1 for w in query_words if w in text_lower)
    return matches / max(len(query_words), 1)


def _tfidf_search(query: str, chunks: list, top_k: int = 5) -> list:
    """Search chunks using TF-IDF index with fallback to keyword search."""
    index = get_index()
    chunk_map = {c.id: c for c in chunks}

    # Try TF-IDF search first
    if index.docs:
        results = index.search(query, top_k=top_k)
        scored = [(score, chunk_map[cid]) for cid, score in results if cid in chunk_map and score > 0]
        if scored:
            return scored

    # Fallback to keyword search
    scored = []
    for chunk in chunks:
        score = _simple_search_score(query, chunk.content)
        if score > 0:
            scored.append((score, chunk))
    scored.sort(key=lambda x: x[0], reverse=True)
    return scored[:top_k]


# ── Document Routes ────────────────────────────────────

@router.get("/api/documents", response_model=list[DocumentResponse])
def list_documents(
    folder_id: Optional[int] = Query(None),
    tag_id: Optional[int] = Query(None),
    search: Optional[str] = Query(None),
    sort: Optional[str] = Query("created_at"),
    view: Optional[str] = Query("list"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(Document).filter(Document.user_id == user.id)

    if folder_id is not None:
        q = q.filter(Document.folder_id == folder_id)
    if tag_id:
        q = q.filter(Document.tags.any(Tag.id == tag_id))
    if search:
        q = q.filter(Document.title.ilike(f"%{search}%"))

    sort_col = getattr(Document, sort, Document.created_at)
    q = q.order_by(sort_col.desc())

    docs = q.all()
    result = []
    for d in docs:
        tags = [{"id": t.id, "name": t.name, "color": t.color} for t in d.tags]
        result.append(DocumentResponse(
            id=d.id, user_id=d.user_id, title=d.title, filename=d.filename,
            file_type=d.file_type, file_size=d.file_size, status=d.status,
            folder_id=d.folder_id, tags=tags,
            created_at=d.created_at, updated_at=d.updated_at,
        ))
    return result


@router.post("/api/documents", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    folder_id: Optional[int] = Form(None),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # Validate file
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail=f"File type {ext} not supported")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    # Save file
    user_dir = os.path.join(UPLOAD_DIR, str(user.id))
    os.makedirs(user_dir, exist_ok=True)
    file_path = os.path.join(user_dir, file.filename or "upload")
    with open(file_path, "wb") as f:
        f.write(content)

    # Extract text
    text = _extract_text(file.filename or "", content)
    status = "ready" if text and not text.startswith("[") else "processing"

    # Create document record
    doc = Document(
        user_id=user.id,
        folder_id=folder_id,
        title=os.path.splitext(file.filename or "document")[0],
        filename=file.filename or "document",
        file_type=ext.lstrip("."),
        file_size=len(content),
        status=status,
        content_text=text[:50000],  # Limit stored text
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Create chunks for RAG + add to search index
    index = get_index()
    if text and not text.startswith("["):
        chunks = _chunk_text(text)
        for i, chunk in enumerate(chunks):
            new_chunk = DocumentChunk(
                document_id=doc.id, chunk_index=i, content=chunk,
                page_number=None,
            )
            db.add(new_chunk)
            db.flush()  # get the chunk id
            index.add(new_chunk.id, chunk)
        db.commit()

    # Auto-tag document (background — don't block response)
    existing_tag_names = [t.name for t in db.query(Tag).filter(Tag.user_id == user.id).all()]
    try:
        import asyncio
        loop = asyncio.get_event_loop()
        suggested = loop.run_until_complete(auto_tag_document(text[:3000], existing_tag_names))
        for tag_name in suggested[:5]:
            tag = db.query(Tag).filter(Tag.name == tag_name, Tag.user_id == user.id).first()
            if not tag:
                tag = Tag(name=tag_name, color="#6d5cff", user_id=user.id)
                db.add(tag)
                db.flush()
            db.add(DocumentTag(document_id=doc.id, tag_id=tag.id))
        db.commit()
    except Exception:
        pass  # auto-tagging is best-effort

    return DocumentResponse(
        id=doc.id, user_id=doc.user_id, title=doc.title, filename=doc.filename,
        file_type=doc.file_type, file_size=doc.file_size, status=doc.status,
        folder_id=doc.folder_id, tags=[],
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


@router.get("/api/documents/{doc_id}", response_model=DocumentResponse)
def get_document(
    doc_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    tags = [{"id": t.id, "name": t.name, "color": t.color} for t in doc.tags]
    return DocumentResponse(
        id=doc.id, user_id=doc.user_id, title=doc.title, filename=doc.filename,
        file_type=doc.file_type, file_size=doc.file_size, status=doc.status,
        folder_id=doc.folder_id, tags=tags,
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


@router.delete("/api/documents/{doc_id}")
def delete_document(
    doc_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove chunks from search index
    index = get_index()
    chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id).all()
    for chunk in chunks:
        index.remove(chunk.id)

    # Delete file from disk
    file_path = os.path.join(UPLOAD_DIR, str(user.id), doc.filename)
    if os.path.exists(file_path):
        os.remove(file_path)

    db.delete(doc)
    db.commit()
    return {"message": "Document deleted"}


@router.put("/api/documents/{doc_id}/move", response_model=DocumentResponse)
def move_document(
    doc_id: int,
    req: DocumentMoveRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if req.folder_id:
        folder = db.query(Folder).filter(Folder.id == req.folder_id, Folder.user_id == user.id).first()
        if not folder:
            raise HTTPException(status_code=404, detail="Folder not found")

    doc.folder_id = req.folder_id
    db.commit()
    db.refresh(doc)
    tags = [{"id": t.id, "name": t.name, "color": t.color} for t in doc.tags]
    return DocumentResponse(
        id=doc.id, user_id=doc.user_id, title=doc.title, filename=doc.filename,
        file_type=doc.file_type, file_size=doc.file_size, status=doc.status,
        folder_id=doc.folder_id, tags=tags,
        created_at=doc.created_at, updated_at=doc.updated_at,
    )


@router.get("/api/documents/{doc_id}/chunks")
def get_chunks(
    doc_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    chunks = db.query(DocumentChunk).filter(
        DocumentChunk.document_id == doc_id
    ).order_by(DocumentChunk.chunk_index).all()

    return [
        {"id": c.id, "chunk_index": c.chunk_index, "content": c.content, "page_number": c.page_number}
        for c in chunks
    ]


# ── Document Q&A (RAG) ────────────────────────────────

@router.post("/api/documents/qa", response_model=DocumentQAResponse)
def document_qa(
    req: DocumentQARequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if req.document_id:
        # Single document Q&A
        doc = db.query(Document).filter(
            Document.id == req.document_id, Document.user_id == user.id
        ).first()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        chunks = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == req.document_id
        ).all()
    elif req.cross_document:
        # Cross-document Q&A
        chunks = db.query(DocumentChunk).join(Document).filter(
            Document.user_id == user.id
        ).all()
    else:
        raise HTTPException(status_code=400, detail="Either document_id or cross_document required")

    if not chunks:
        return DocumentQAResponse(answer="No document content found to answer from.", sources=[])

    # TF-IDF semantic search with keyword fallback
    scored = _tfidf_search(req.question, chunks, top_k=5)
    top_chunks = scored if scored else [(0.1, chunks[0])]

    # Build context
    context = "\n\n".join([c.content[:500] for _, c in top_chunks])
    answer = f"Based on the document content:\n\n{context[:2000]}\n\n---\n*Note: This is a basic text search. For AI-powered answers, configure an AI provider.*"

    sources = []
    for score, chunk in top_chunks:
        doc = db.query(Document).filter(Document.id == chunk.document_id).first()
        sources.append({
            "document_id": chunk.document_id,
            "document_title": doc.title if doc else "Unknown",
            "chunk_index": chunk.chunk_index,
            "relevance": round(score, 2),
            "preview": chunk.content[:200],
        })

    # Save Q&A record
    qa = DocumentQA(
        document_id=req.document_id,
        user_id=user.id,
        question=req.question,
        answer=answer,
        sources=sources,
        cross_document=req.cross_document,
    )
    db.add(qa)
    db.commit()

    return DocumentQAResponse(answer=answer, sources=sources)


@router.post("/api/documents/{doc_id}/summarize", response_model=DocumentSummarizeResponse)
def summarize_document(
    doc_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    text = doc.content_text or ""
    if not text:
        return DocumentSummarizeResponse(
            summary="No content available to summarize.",
            key_points=[]
        )

    # Simple extractive summary
    sentences = re.split(r'(?<=[.!?])\s+', text)
    summary = " ".join(sentences[:5]) if sentences else text[:500]

    # Extract key points (sentences with certain keywords)
    keywords = ["important", "key", "main", "significant", "note", "must", "should", "require"]
    key_points = []
    for s in sentences:
        if any(kw in s.lower() for kw in keywords):
            key_points.append(s.strip())
        if len(key_points) >= 5:
            break

    if not key_points:
        key_points = [s.strip() for s in sentences[:3] if len(s.strip()) > 20]

    return DocumentSummarizeResponse(summary=summary, key_points=key_points)


# ── Folder Routes ──────────────────────────────────────

@router.get("/api/folders", response_model=list[FolderResponse])
def list_folders(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    folders = db.query(Folder).filter(Folder.user_id == user.id).all()
    result = []
    for f in folders:
        count = db.query(Document).filter(Document.folder_id == f.id).count()
        result.append(FolderResponse(
            id=f.id, name=f.name, color=f.color, document_count=count
        ))
    return result


@router.post("/api/folders", response_model=FolderResponse)
def create_folder(
    req: FolderCreate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    folder = Folder(user_id=user.id, name=req.name, color=req.color)
    db.add(folder)
    db.commit()
    db.refresh(folder)
    return FolderResponse(id=folder.id, name=folder.name, color=folder.color, document_count=0)


@router.delete("/api/folders/{folder_id}")
def delete_folder(
    folder_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    folder = db.query(Folder).filter(Folder.id == folder_id, Folder.user_id == user.id).first()
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")

    # Move documents out of folder
    db.query(Document).filter(Document.folder_id == folder_id).update({"folder_id": None})
    db.delete(folder)
    db.commit()
    return {"message": "Folder deleted"}


# ── Tag Routes ─────────────────────────────────────────

@router.get("/api/tags", response_model=list[TagResponse])
def list_tags(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tags = db.query(Tag).filter(Tag.user_id == user.id).all()
    return [TagResponse(id=t.id, name=t.name, color=t.color) for t in tags]


@router.post("/api/tags", response_model=TagResponse)
def create_tag(
    req: TagCreate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tag = Tag(user_id=user.id, name=req.name, color=req.color)
    db.add(tag)
    db.commit()
    db.refresh(tag)
    return TagResponse(id=tag.id, name=tag.name, color=tag.color)


@router.delete("/api/tags/{tag_id}")
def delete_tag(
    tag_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tag = db.query(Tag).filter(Tag.id == tag_id, Tag.user_id == user.id).first()
    if not tag:
        raise HTTPException(status_code=404, detail="Tag not found")
    db.delete(tag)
    db.commit()
    return {"message": "Tag deleted"}


@router.post("/api/documents/{doc_id}/tags")
def assign_tag(
    doc_id: int,
    req: TagAssignRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    tag = db.query(Tag).filter(Tag.id == req.tag_id, Tag.user_id == user.id).first()
    if not tag:
        raise HTTPException(status_code=404, detail="Tag not found")

    existing = db.query(DocumentTag).filter(
        DocumentTag.document_id == doc_id, DocumentTag.tag_id == req.tag_id
    ).first()
    if not existing:
        db.add(DocumentTag(document_id=doc_id, tag_id=req.tag_id))
        db.commit()

    return {"message": "Tag assigned"}


@router.delete("/api/documents/{doc_id}/tags/{tag_id}")
def remove_tag(
    doc_id: int,
    tag_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    db.query(DocumentTag).filter(
        DocumentTag.document_id == doc_id, DocumentTag.tag_id == tag_id
    ).delete()
    db.commit()
    return {"message": "Tag removed"}


# ── Document Sharing ───────────────────────────────────

class DocumentShareRequest(BaseModel):
    email: str
    permission: str = "view"  # view or edit


class DocumentShareResponse(BaseModel):
    id: int
    document_id: int
    shared_with_email: str
    permission: str
    created_at: datetime
    model_config = {"from_attributes": True}


@router.post("/api/documents/{doc_id}/share", response_model=DocumentShareResponse)
def share_document(
    doc_id: int,
    req: DocumentShareRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Share a document with another user by email."""
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if req.permission not in ("view", "edit"):
        raise HTTPException(status_code=400, detail="Permission must be 'view' or 'edit'")

    # Check for existing share with same email
    existing = db.query(DocumentShare).filter(
        DocumentShare.document_id == doc_id,
        DocumentShare.shared_with_email == req.email,
    ).first()
    if existing:
        existing.permission = req.permission
        db.commit()
        db.refresh(existing)
        return DocumentShareResponse(
            id=existing.id,
            document_id=existing.document_id,
            shared_with_email=existing.shared_with_email,
            permission=existing.permission,
            created_at=existing.created_at,
        )

    share = DocumentShare(
        document_id=doc_id,
        shared_by_user_id=user.id,
        shared_with_email=req.email,
        permission=req.permission,
    )
    db.add(share)
    db.commit()
    db.refresh(share)
    return DocumentShareResponse(
        id=share.id,
        document_id=share.document_id,
        shared_with_email=share.shared_with_email,
        permission=share.permission,
        created_at=share.created_at,
    )


@router.get("/api/documents/shared-with-me", response_model=list[DocumentResponse])
def shared_with_me(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List documents that other users have shared with the current user."""
    shares = db.query(DocumentShare).filter(
        DocumentShare.shared_with_email == user.email,
    ).all()

    result = []
    for share in shares:
        doc = db.query(Document).filter(Document.id == share.document_id).first()
        if doc:
            tags = [{"id": t.id, "name": t.name, "color": t.color} for t in doc.tags]
            result.append(DocumentResponse(
                id=doc.id, user_id=doc.user_id, title=doc.title, filename=doc.filename,
                file_type=doc.file_type, file_size=doc.file_size, status=doc.status,
                folder_id=doc.folder_id, tags=tags,
                created_at=doc.created_at, updated_at=doc.updated_at,
            ))
    return result


@router.delete("/api/documents/{doc_id}/shares/{share_id}")
def remove_share(
    doc_id: int,
    share_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Remove a document share."""
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    share = db.query(DocumentShare).filter(
        DocumentShare.id == share_id,
        DocumentShare.document_id == doc_id,
    ).first()
    if not share:
        raise HTTPException(status_code=404, detail="Share not found")

    db.delete(share)
    db.commit()
    return {"message": "Share removed"}
